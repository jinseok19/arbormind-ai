"""리포트 생성 모듈 (PDF + Word)"""

from datetime import datetime
from typing import Dict, Optional
from pathlib import Path
import io
import matplotlib.pyplot as plt
import matplotlib
import numpy as np

# PDF 생성용
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Word 생성용
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    """PDF 및 Word 리포트 생성 클래스"""
    
    def __init__(self):
        self.output_dir = Path("reports")
        self.pdf_dir = self.output_dir / "pdf"
        self.word_dir = self.output_dir / "word"
        self.charts_dir = self.output_dir / "charts"
        
        # 디렉토리 생성
        self.pdf_dir.mkdir(parents=True, exist_ok=True)
        self.word_dir.mkdir(parents=True, exist_ok=True)
        self.charts_dir.mkdir(parents=True, exist_ok=True)
        
        # 한글 폰트 등록 (Windows)
        try:
            # 맑은 고딕
            pdfmetrics.registerFont(TTFont('Malgun', 'malgun.ttf'))
            pdfmetrics.registerFont(TTFont('MalgunBd', 'malgunbd.ttf'))
            self.font_name = 'Malgun'
            self.font_name_bold = 'MalgunBd'
        except:
            try:
                # 나눔고딕 (대안)
                pdfmetrics.registerFont(TTFont('NanumGothic', 'NanumGothic.ttf'))
                self.font_name = 'NanumGothic'
                self.font_name_bold = 'NanumGothic'
            except:
                # 기본 폰트 사용
                self.font_name = 'Helvetica'
                self.font_name_bold = 'Helvetica-Bold'
        
        # matplotlib 한글 폰트 설정
        matplotlib.rcParams['font.family'] = 'Malgun Gothic'
        matplotlib.rcParams['axes.unicode_minus'] = False
    
    def generate_pdf(
        self,
        analysis_id: str,
        park_info: Dict,
        areas: Dict,
        carbon: Dict,
        original_image_path: Optional[str] = None,
        overlay_image_path: Optional[str] = None
    ) -> str:
        """
        PDF 리포트 생성
        
        Args:
            analysis_id: 분석 ID
            park_info: 공원 정보
            areas: 면적 데이터
            carbon: 탄소 데이터
            overlay_image_path: 오버레이 이미지 경로
        
        Returns:
            생성된 PDF 파일 경로
        """
        # 파일명
        filename = f"{analysis_id}_report.pdf"
        filepath = self.pdf_dir / filename
        
        # PDF 문서 생성
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            topMargin=2*cm,
            bottomMargin=2*cm,
            leftMargin=2*cm,
            rightMargin=2*cm
        )
        
        # 차트 생성
        from utils.chart_generator import ChartGenerator
        chart_gen = ChartGenerator()
        
        pie_chart_path = self.charts_dir / f"{analysis_id}_pie.png"
        bar_chart_path = self.charts_dir / f"{analysis_id}_bar.png"
        carbon_chart_path = self.charts_dir / f"{analysis_id}_carbon.png"
        
        chart_gen.create_professional_pie_chart(areas, str(pie_chart_path))
        chart_gen.create_professional_bar_chart(areas, str(bar_chart_path))
        if carbon and carbon.get('total_tco2_yr'):
            chart_gen.create_carbon_chart(carbon, str(carbon_chart_path))
        
        # 한글 스타일 생성
        styles = getSampleStyleSheet()
        
        # 제목 스타일
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=self.font_name_bold,
            fontSize=24,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # 헤딩 스타일
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=self.font_name_bold,
            fontSize=14,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # 일반 텍스트 스타일
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=self.font_name,
            fontSize=10,
            leading=14
        )
        
        # 컨텐츠 빌드
        story = []
        
        # 제목
        story.append(Paragraph("ArborMind AI 공원 탄소흡수 추정 리포트", title_style))
        story.append(Spacer(1, 0.3*cm))
        
        # 부제
        subtitle = ParagraphStyle(
            'Subtitle',
            parent=normal_style,
            fontSize=12,
            textColor=colors.HexColor('#7F8C8D'),
            alignment=TA_CENTER
        )
        story.append(Paragraph("AI 기반 식생 분석 및 탄소흡수량 평가 보고서", subtitle))
        story.append(Spacer(1, 0.8*cm))
        
        # 기본 정보
        story.append(Paragraph("1. 기본 정보", heading_style))
        info_data = [
            ["공원명", park_info.get('name', '-')],
            ["위치", park_info.get('location', '-')],
            ["분석일", datetime.now().strftime('%Y-%m-%d %H:%M')],
            ["총 면적", f"{park_info.get('total_area_m2', 0):,.0f} ㎡" if park_info.get('total_area_m2') else "미입력"]
        ]
        info_table = Table(info_data, colWidths=[4*cm, 12*cm])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ECF0F1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#2C3E50')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), self.font_name_bold),
            ('FONTNAME', (1, 0), (1, -1), self.font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.8*cm))
        
        # 식생 타입별 면적
        story.append(Paragraph("2. 식생 타입별 면적 요약", heading_style))
        
        type_labels = {
            'BUILDING': '건물',
            'ROAD': '도로',
            'WATER': '물',
            'FOREST': '숲',
            'TREE': '나무',
            'GRASS': '초지',
            'WETLAND': '습지',
            'SOIL': '토양'
        }
        
        area_data = [["타입", "면적(㎡)", "비율", "탄소흡수량"]]
        for veg_type, data in areas.items():
            label = type_labels.get(veg_type, veg_type)
            area_m2 = f"{data['area_m2']:,.1f}" if data['area_m2'] else "-"
            ratio = f"{data['ratio_percent']:.1f}%"
            carbon_val = f"{carbon['by_type'].get(veg_type, 0):.2f} tCO₂" if carbon['by_type'] else "-"
            area_data.append([label, area_m2, ratio, carbon_val])
        
        area_table = Table(area_data, colWidths=[4*cm, 4*cm, 3*cm, 5*cm])
        area_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495E')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), self.font_name_bold),
            ('FONTNAME', (0, 1), (-1, -1), self.font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')])
        ]))
        story.append(area_table)
        story.append(Spacer(1, 0.8*cm))
        
        # 탄소흡수량
        if carbon.get('total_tco2_yr'):
            story.append(Paragraph("3. 연간 탄소흡수량 총계", heading_style))
            
            # 강조 박스
            carbon_box_style = ParagraphStyle(
                'CarbonBox',
                parent=normal_style,
                fontSize=16,
                textColor=colors.HexColor('#27AE60'),
                alignment=TA_CENTER,
                leading=22
            )
            carbon_text = f"<b>총 탄소흡수량: {carbon['total_tco2_yr']:.2f} tCO₂/yr</b>"
            
            # 박스 테이블로 강조
            carbon_data = [[Paragraph(carbon_text, carbon_box_style)]]
            carbon_table = Table(carbon_data, colWidths=[16*cm])
            carbon_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#E8F8F5')),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#27AE60')),
                ('TOPPADDING', (0, 0), (-1, -1), 15),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
            ]))
            story.append(carbon_table)
            story.append(Spacer(1, 0.5*cm))
            
            # 탄소 차트 추가
            if carbon_chart_path.exists():
                story.append(Paragraph("탄소흡수량 비교", normal_style))
                story.append(Spacer(1, 0.3*cm))
                carbon_img = RLImage(str(carbon_chart_path), width=14*cm, height=8*cm)
                story.append(carbon_img)
                story.append(Spacer(1, 0.5*cm))
        
        # 이미지 (있을 경우)
        story.append(Paragraph("4. 항공 사진 분석 이미지", heading_style))
        
        # 원본 이미지
        if original_image_path and Path(original_image_path).exists():
            story.append(Paragraph("4-1. 원본 이미지", normal_style))
            try:
                img = RLImage(original_image_path)
                img_width = 14*cm
                aspect = img.imageHeight / img.imageWidth
                img_height = img_width * aspect
                if img_height > 10*cm:
                    img_height = 10*cm
                    img_width = img_height / aspect
                
                img = RLImage(original_image_path, width=img_width, height=img_height)
                story.append(img)
                story.append(Spacer(1, 0.5*cm))
            except Exception as e:
                story.append(Paragraph(f"원본 이미지 로드 실패: {str(e)}", normal_style))
                story.append(Spacer(1, 0.3*cm))
        
        # 세그멘테이션 결과 이미지
        if overlay_image_path and Path(overlay_image_path).exists():
            story.append(Paragraph("4-2. 세그멘테이션 분석 결과", normal_style))
            try:
                img = RLImage(overlay_image_path)
                img_width = 14*cm
                aspect = img.imageHeight / img.imageWidth
                img_height = img_width * aspect
                if img_height > 12*cm:
                    img_height = 12*cm
                    img_width = img_height / aspect
                
                img = RLImage(overlay_image_path, width=img_width, height=img_height)
                story.append(img)
                story.append(Spacer(1, 0.5*cm))
            except Exception as e:
                story.append(Paragraph(f"세그멘테이션 이미지 로드 실패: {str(e)}", normal_style))
                story.append(Spacer(1, 0.5*cm))
        
        # 페이지 나누기
        story.append(PageBreak())
        
        # 시각화 차트 추가
        story.append(Paragraph("5. 면적 비율 시각화", heading_style))
        story.append(Spacer(1, 0.5*cm))
        
        if pie_chart_path.exists() and bar_chart_path.exists():
            # 파이 차트
            story.append(Paragraph("5-1. 파이 차트", normal_style))
            story.append(Spacer(1, 0.3*cm))
            pie_img = RLImage(str(pie_chart_path), width=12*cm, height=9*cm)
            story.append(pie_img)
            story.append(Spacer(1, 0.5*cm))
            
            # 막대 차트
            story.append(Paragraph("5-2. 막대 차트", normal_style))
            story.append(Spacer(1, 0.3*cm))
            bar_img = RLImage(str(bar_chart_path), width=12*cm, height=9*cm)
            story.append(bar_img)
            story.append(Spacer(1, 0.5*cm))
        
        # 산정 방법
        story.append(Paragraph("6. 산정 방법", heading_style))
        method_text = """
        본 결과는 AI 기반 이미지 분석을 통해 식생 타입을 공간 분해하여 면적을 산출하고,
        타입별 대표 탄소흡수 계수를 적용해 연간 탄소흡수량을 추정한 값입니다.
        <br/><br/>
        <b>분석 방법:</b><br/>
        1. 항공/드론 이미지 전처리 및 정규화<br/>
        2. AI 기반 세그멘테이션을 통한 8가지 타입 분류<br/>
        3. 픽셀 단위 면적 계산 및 실제 면적 변환<br/>
        4. 타입별 탄소흡수 계수 적용 및 총량 산정
        """
        story.append(Paragraph(method_text, normal_style))
        story.append(Spacer(1, 0.5*cm))
        
        # 한계
        story.append(Paragraph("7. 데이터 품질 및 한계", heading_style))
        limit_text = """
        <b>현재 단계:</b> MVP 프로토타입 (컬러 기반 분류)<br/><br/>
        본 수치는 대표값 기반 추정치이며, 다음 단계의 고도화를 통해 정확도가 개선됩니다:<br/>
        • 딥러닝 기반 정밀 세그멘테이션 모델 적용<br/>
        • 수종별 세부 분류 및 개별 탄소계수 적용<br/>
        • 수령, 생육 상태, 지역 특성 반영<br/>
        • 실측 데이터 기반 검증 및 보정
        """
        story.append(Paragraph(limit_text, normal_style))
        story.append(Spacer(1, 0.5*cm))
        
        # 푸터
        footer_style = ParagraphStyle(
            'Footer',
            parent=normal_style,
            fontSize=8,
            textColor=colors.HexColor('#95A5A6'),
            alignment=TA_CENTER
        )
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}", footer_style))
        story.append(Paragraph("Powered by ArborMind AI v1.0", footer_style))
        
        # PDF 생성
        doc.build(story)
        
        return str(filepath)
    
    def generate_word(
        self,
        analysis_id: str,
        park_info: Dict,
        areas: Dict,
        carbon: Dict,
        original_image_path: Optional[str] = None,
        overlay_image_path: Optional[str] = None
    ) -> str:
        """
        Word 리포트 생성
        
        Args:
            analysis_id: 분석 ID
            park_info: 공원 정보
            areas: 면적 데이터
            carbon: 탄소 데이터
            overlay_image_path: 오버레이 이미지 경로
        
        Returns:
            생성된 Word 파일 경로
        """
        # 파일명
        filename = f"{analysis_id}_report.docx"
        filepath = self.word_dir / filename
        
        # 문서 생성
        doc = Document()
        
        # 제목
        title = doc.add_heading('ArborMind AI 공원 탄소흡수 추정 리포트', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 기본 정보
        doc.add_heading('1. 기본 정보', level=2)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = '공원명'
        table.rows[0].cells[1].text = park_info.get('name', '-')
        table.rows[1].cells[0].text = '위치'
        table.rows[1].cells[1].text = park_info.get('location', '-')
        table.rows[2].cells[0].text = '분석일'
        table.rows[2].cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M')
        table.rows[3].cells[0].text = '총 면적'
        table.rows[3].cells[1].text = f"{park_info.get('total_area_m2', 0):,.0f} ㎡" if park_info.get('total_area_m2') else "미입력"
        
        doc.add_paragraph()
        
        # 식생 타입별 면적
        doc.add_heading('2. 식생 타입별 면적 요약', level=2)
        
        type_labels = {
            'BUILDING': '건물',
            'ROAD': '도로',
            'WATER': '물',
            'FOREST': '숲',
            'TREE': '나무',
            'GRASS': '초지',
            'WETLAND': '습지',
            'SOIL': '토양'
        }
        
        table = doc.add_table(rows=len(areas)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # 헤더
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '타입'
        hdr_cells[1].text = '면적(㎡)'
        hdr_cells[2].text = '비율'
        hdr_cells[3].text = '탄소흡수량'
        
        # 데이터
        for i, (veg_type, data) in enumerate(areas.items(), start=1):
            label = type_labels.get(veg_type, veg_type)
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = f"{data['area_m2']:,.1f}" if data['area_m2'] else "-"
            row_cells[2].text = f"{data['ratio_percent']:.1f}%"
            carbon_val = carbon['by_type'].get(veg_type, 0) if carbon['by_type'] else 0
            row_cells[3].text = f"{carbon_val:.2f} tCO₂" if carbon_val else "-"
        
        doc.add_paragraph()
        
        # 탄소흡수량
        if carbon.get('total_tco2_yr'):
            doc.add_heading('3. 연간 탄소흡수량 총계', level=2)
            p = doc.add_paragraph()
            run = p.add_run(f"총 탄소흡수량: {carbon['total_tco2_yr']:.2f} tCO₂/yr")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 128, 0)
            
            doc.add_paragraph()
        
        # 이미지 (있을 경우)
        doc.add_heading('4. 항공 사진 분석 이미지', level=2)
        
        # 원본 이미지
        if original_image_path and Path(original_image_path).exists():
            doc.add_heading('4-1. 원본 이미지', level=3)
            try:
                doc.add_picture(original_image_path, width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"원본 이미지 로드 실패: {str(e)}")
            doc.add_paragraph()
        
        # 세그멘테이션 결과
        if overlay_image_path and Path(overlay_image_path).exists():
            doc.add_heading('4-2. 세그멘테이션 분석 결과', level=3)
            try:
                doc.add_picture(overlay_image_path, width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f"세그멘테이션 이미지 로드 실패: {str(e)}")
            doc.add_paragraph()
        
        # 산정 방법
        doc.add_heading('5. 산정 방법', level=2)
        doc.add_paragraph(
            "본 결과는 입력 이미지 기반 식생 타입을 공간 분해하여 면적을 산출하고, "
            "타입별 대표 계수를 적용해 연간 탄소흡수량을 추정한 값입니다."
        )
        
        # 한계
        doc.add_heading('6. 한계 및 고도화 방향', level=2)
        doc.add_paragraph(
            "본 수치는 MVP 단계의 대표값 기반 추정치이며, "
            "향후 수종/생육/지역/실측 데이터 결합 시 정확도가 개선됩니다."
        )
        
        # 문서 저장
        doc.save(str(filepath))
        
        return str(filepath)

